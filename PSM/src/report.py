from __future__ import annotations

import io
from typing import Any

import pandas as pd

from src.risk import METRIC_NAMES
from src.sentiment import sentiment_distribution
from src.stance import stance_distribution
from src.timeline import aggregate_by_time, detect_peak_period, estimate_lifecycle_stage
from src.uncertainty import uncertainty_summary


def _format_ratio(df: pd.DataFrame, label_col: str, label: str) -> str:
    if label_col not in df.columns or df.empty:
        return "0%"
    return f"{(df[label_col] == label).mean() * 100:.2f}%"


def _top_keywords(keyword_df: pd.DataFrame | None, top_n: int = 10) -> str:
    if keyword_df is None or keyword_df.empty or "word" not in keyword_df.columns:
        return "暂无关键词结果"
    return "、".join(keyword_df["word"].head(top_n).astype(str).tolist())


def build_report_context(
    df: pd.DataFrame,
    quality_report: dict[str, Any] | None = None,
    keyword_df: pd.DataFrame | None = None,
    cluster_keywords: dict[int, list[str]] | None = None,
    risk_result: dict[str, Any] | None = None,
    target: str = "",
    source_note: str = "",
) -> dict[str, Any]:
    source_col = "clean_content" if "clean_content" in df.columns else "content"
    time_series = aggregate_by_time(df)
    uncertainty = uncertainty_summary(df)
    cluster_names = []
    for cluster_id, words in (cluster_keywords or {}).items():
        cluster_names.append(f"主题{cluster_id}（{'、'.join(words[:5])}）")

    examples = []
    if source_col in df.columns:
        examples = df[source_col].dropna().astype(str).head(5).tolist()

    return {
        "total_count": len(df),
        "valid_count": quality_report.get("valid_count", len(df)) if quality_report else len(df),
        "platforms": "、".join(df["platform"].dropna().astype(str).unique().tolist())
        if "platform" in df.columns
        else "未提供",
        "time_range": quality_report.get("time_range", "无有效时间字段") if quality_report else "无有效时间字段",
        "peak_date": detect_peak_period(time_series),
        "lifecycle_stage": estimate_lifecycle_stage(time_series),
        "negative_ratio": _format_ratio(df, "sentiment_label", "negative"),
        "neutral_ratio": _format_ratio(df, "sentiment_label", "neutral"),
        "positive_ratio": _format_ratio(df, "sentiment_label", "positive"),
        "topic_count": len(cluster_keywords or {}),
        "topic_names": "；".join(cluster_names) if cluster_names else "暂无聚类结果",
        "dominant_topic": cluster_names[0] if cluster_names else "暂无聚类结果",
        "keywords": _top_keywords(keyword_df),
        "target": target or "未指定",
        "support_ratio": _format_ratio(df, "stance_label", "support"),
        "oppose_ratio": _format_ratio(df, "stance_label", "oppose"),
        "neutral_stance_ratio": _format_ratio(df, "stance_label", "neutral"),
        "uncertainty_ratio": f"{uncertainty['ratio']:.2f}%",
        "uncertainty_terms": "、".join(uncertainty["terms"]) if uncertainty["terms"] else "未发现高频不确定表达",
        "risk_score": risk_result.get("score", 0) if risk_result else 0,
        "risk_level": risk_result.get("level", "未计算") if risk_result else "未计算",
        "risk_reasons": "；".join(risk_result.get("reasons", [])) if risk_result else "未计算风险来源",
        "representative_comments": "\n".join(f"- {item}" for item in examples) if examples else "- 暂无代表评论",
        "recommendations": generate_recommendations(risk_result.get("level", "") if risk_result else ""),
        "source_note": source_note,
    }


def generate_recommendations(risk_level: str) -> str:
    if risk_level in {"较高风险", "高风险"}:
        return "建议优先核查高频争议点，及时补充权威信息，并对不确定表达较多的内容进行澄清。"
    if risk_level == "中风险":
        return "建议持续观察情绪变化和主题集中度，针对核心诉求准备解释材料。"
    return "当前文本信号风险较低，建议保留常规监测，并关注后续是否出现新的争议主题。"


def generate_markdown_report(context: dict[str, Any]) -> str:
    return f"""# 舆情分析报告

## 一、数据来源与样本说明
本次分析共导入 {context['total_count']} 条文本，清洗后保留 {context['valid_count']} 条有效文本。数据来源包括 {context['platforms']}，时间范围为 {context['time_range']}。

{context['source_note']}

## 二、事件热度变化
从时间分布看，讨论量在 {context['peak_date']} 达到峰值。系统初步判断该事件处于 {context['lifecycle_stage']}。

## 三、情感倾向分析
总体情感分布显示，负向文本占比为 {context['negative_ratio']}，中性文本占比为 {context['neutral_ratio']}，正向文本占比为 {context['positive_ratio']}。

## 四、关键词、主题与观点结构
本次文本中的核心关键词包括：{context['keywords']}。

系统识别出 {context['topic_count']} 类主要讨论主题：{context['topic_names']}。其中占比较高的是 {context['dominant_topic']}。

## 五、立场分析
针对目标对象“{context['target']}”，支持立场占比为 {context['support_ratio']}，反对立场占比为 {context['oppose_ratio']}，中立立场占比为 {context['neutral_stance_ratio']}。

## 六、不确定信息提示
系统识别到不确定信息表达占比为 {context['uncertainty_ratio']}，相关高频表达包括：{context['uncertainty_terms']}。

## 七、风险辅助研判
系统给出的舆情风险评分为 {context['risk_score']}，风险等级为 {context['risk_level']}。主要风险来源包括：{context['risk_reasons']}。

本风险等级为基于文本信号的辅助研判结果，不代表对现实事件风险的确定性判断。

## 八、代表性评论
{context['representative_comments']}

## 九、建议
{context['recommendations']}

## 十、方法局限
本报告基于用户上传数据和文本分析模型自动生成，结果受样本来源、文本质量、模型误差和语境差异影响，仅用于辅助分析。第一版采用轻量词典、TF-IDF 与 K-means 等可解释方法，后续可接入 SBERT、BERTopic 与 BERT 情感模型提升语义理解能力。

## 十一、文献依据
本系统的方法设计参考了网络舆情情感演化、LDA 主题模型、BERTopic、SBERT、社交媒体事件摘要、社会媒体立场检测、社会网络谣言检测和突发事件群体极化风险评估等研究。当前版本将相关研究转化为可运行的轻量分析链路：数据清洗、关键词提取、观点聚类、情感分析、目标对象立场检测、不确定信息识别、风险辅助评分和结构化报告生成。
"""


def generate_docx_report_bytes(markdown_report: str) -> bytes:
    from docx import Document

    document = Document()
    for line in markdown_report.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
